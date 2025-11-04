import fitz # PyMuPDF library for PDF processing
import docx # python-docx library for DOCX processing
import os
import tempfile # For creating temporary files securely
import json
import requests # For making HTTP requests (will still be used, but to Google API now)
# Removed sklearn and spacy imports as matching logic will move to Gemini
import logging # For logging application events and errors
import google.generativeai as genai
from fastapi import FastAPI, File, UploadFile, Form, HTTPException, BackgroundTasks,Query # FastAPI core components
from fastapi.middleware.cors import CORSMiddleware # For handling Cross-Origin Resource Sharing
from fastapi.responses import FileResponse # For sending files as responses
from pydantic import BaseModel # For data validation with request bodies
from typing import List, Optional # For type hinting
import uvicorn # ASGI server for running FastAPI
from job_apis import fetch_jobs
from dotenv import load_dotenv
load_dotenv()
# --- Logging Configuration ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- FastAPI Application Initialization ---
app = FastAPI(
    title="Resume Matcher & Generator API",
    description="API for matching resumes against job descriptions and generating/improving resumes using AI."
)
# --- Google Gemini API Integration ---
# Replace YOUR_GOOGLE_GEMINI_API_KEY with your actual Google Gemini API Key
 # Your API key 
GOOGLE_GEMINI_API_KEY = os.getenv("GOOGLE_GEMINI_API_KEY")
GEMINI_MODEL = "gemini-2.0-flash"

# CRITICAL FIX: Define the URL as a literal string to prevent any markdown parsing issues.
# Ensure no markdown brackets or parentheses are introduced here.
GEMINI_API_BASE_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"
logging.info(f"DEBUG: GEMINI_API_BASE_URL (global) is set to: '{GOOGLE_GEMINI_API_KEY}'")

# --- CORS (Cross-Origin Resource Sharing) Setup ---
# This allows your frontend application (e.g., React app running on localhost:3000)
# to make requests to this backend API.
# FIX: Ensure origins are clean string literals
origins = ["http://localhost:3000", "http://127.0.0.1:3000"]
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],     # Allow all origins (frontend URLs)
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# --- USER AUTHENTICATION START ---
# users_db = {} # This is a simple in-memory store. For production, use a proper database.

# class User(BaseModel):
#     email: str
#     password: str
#     name: str = None

# @app.post("/signup")
# def signup(user: User):
#     if user.email in users_db:
#         raise HTTPException(status_code=400, detail="User already exists")
#     users_db[user.email] = user.password # Store password (in a real app, hash this!)
#     logging.info(f"User {user.email} signed up.")
#     return {"success": True, "message": "User registered successfully."}

# @app.post("/login")
# def login(user: User):
#     if user.email not in users_db or users_db[user.email] != user.password:
#         raise HTTPException(status_code=401, detail="Invalid credentials. Please check your email and password.")
#     logging.info(f"User {user.email} logged in.")
#     return {"success": True, "message": "Login successful."}
# # --- USER AUTHENTICATION END ---





@app.get("/jobs")
def get_jobs(query: str = Query(..., description="Job title or skill"),
             location: str = Query("India", description="Job location")):
    jobs = fetch_jobs(query, location)
    if not jobs:
        return {"message": "No jobs found"}
    return {"total": len(jobs), "results": jobs}






# Request body model
class InterviewRequest(BaseModel):
    message: str

@app.post("/interview")
async def interview(req: InterviewRequest):
    try:
        print("in /interview api")

        # ✅ Configure the Gemini API first
        genai.configure(api_key=GOOGLE_GEMINI_API_KEY)

        # Create a prompt for the AI interviewer
        prompt = f"""
        You are an AI interviewer helping a candidate prepare for technical interviews.
        Respond conversationally but professionally.
        The user said: "{req.message}"
        Your reply should be short (2-3 sentences), encouraging, and relevant.
        """

        # Use the Gemini model
        model = genai.GenerativeModel(GEMINI_MODEL)
        response = model.generate_content(prompt)

        ai_reply = response.text.strip() if hasattr(response, "text") else "Sorry, I couldn’t generate a reply."
        return {"reply": ai_reply}

    except Exception as e:
        logging.error(f"Interview route error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# --- Helper Functions for Text Extraction ---

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extracts text content from a PDF file's bytes.
    Raises HTTPException if the PDF file is corrupted or cannot be processed.
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
        return text
    except Exception as e:
        logging.error(f"Error extracting text from PDF: {e}")
        raise HTTPException(status_code=400, detail="Could not process PDF file. It might be corrupted, password-protected, or an image-only PDF.")

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts text content from a DOCX file's bytes.
    Raises HTTPException if the DOCX file is corrupted or cannot be processed.
    Uses a temporary file as python-docx requires a file path.
    """
    temp_path = ""
    try:
        # Create a temporary file to write the bytes to, as python-docx needs a file path
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_f:
            temp_f.write(file_bytes)
            temp_path = temp_f.name
        doc = docx.Document(temp_path)
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        logging.error(f"Error extracting text from DOCX: {e}")
        raise HTTPException(status_code=400, detail="Could not process DOCX file. It might be corrupted or malformed.")
    finally:
        # Ensure the temporary file is removed regardless of success or failure
        if os.path.exists(temp_path):
            os.remove(temp_path)

# --- DOCX Generation from JSON (Helper Function) ---
def generate_resume_from_json(data: dict):
    """
    Creates a DOCX document from a dictionary containing structured resume data.
    This function structures the content parsed from the AI's JSON output into a Word document.
    """
    doc = docx.Document()
    
    # Set basic page margins for a cleaner look
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = docx.shared.Inches(0.75)

    # Add Name and Contact Info (centered)
    if 'name' in data and data['name'].strip():
        name_heading = doc.add_heading(data['name'], 0)
        name_heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    contact_info_parts = []
    contact = data.get('contact', {})
    if 'email' in contact and contact['email'].strip():
        contact_info_parts.append(contact['email'])
    if 'phone' in contact and contact['phone'].strip():
        contact_info_parts.append(contact['phone'])
    if 'linkedin' in contact and contact['linkedin'].strip():
        contact_info_parts.append(contact['linkedin'])
    
    if contact_info_parts:
        contact_para = doc.add_paragraph(' | '.join(contact_info_parts)) # Use a separator for contact info
        contact_para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("") # Add a blank line for spacing

    # Add Summary
    if 'summary' in data and data['summary'].strip():
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(data['summary'])

    # Add Skills
    if 'skills' in data and data['skills']:
        doc.add_heading('Skills', level=1)
        # Display skills as a comma-separated list
        skills_text = ", ".join(data['skills'])
        doc.add_paragraph(skills_text)

    # Add Experience
    if 'experience' in data and data['experience']:
        doc.add_heading('Experience', level=1)
        for exp in data['experience']:
            p = doc.add_paragraph()
            p.add_run(f"{exp.get('job_title', '').strip()}").bold = True # Job Title in Bold
            p.add_run(f" at {exp.get('company', '').strip()}").italic = True # Company in Italic
            p.add_run(f" ({exp.get('dates', '').strip()})").font.size = docx.shared.Pt(10) # Dates smaller
            # Add description as a bullet point
            if exp.get('description', '').strip():
                doc.add_paragraph(exp['description'].strip(), style='ListBullet')

    # Add Education
    if 'education' in data and data['education']:
        doc.add_heading('Education', level=1)
        for edu in data['education']:
            p = doc.add_paragraph()
            p.add_run(f"{edu.get('degree', '').strip()}").bold = True
            p.add_run(f" - {edu.get('institution', '').strip()} ({edu.get('year', '').strip()})")

    # Add Certifications
    if 'certifications' in data and data['certifications']:
        doc.add_heading('Certifications', level=1)
        for cert in data['certifications']:
            doc.add_paragraph(f"- {cert.strip()}", style='ListBullet')

    # Add Projects
    if 'projects' in data and data['projects']:
        doc.add_heading('Projects', level=1)
        for project in data['projects']:
            p = doc.add_paragraph()
            p.add_run(f"{project.get('title', '').strip()}").bold = True
            if project.get('tech_stack'):
                p.add_run(f" - ({', '.join(tech.strip() for tech in project['tech_stack'])})").font.size = docx.shared.Pt(10)
            if project.get('description', '').strip():
                doc.add_paragraph(project['description'].strip(), style='ListBullet')

    return doc




@app.post("/match")
async def match_resume(
    resume: UploadFile = File(..., description="The candidate's resume file (PDF or DOCX format)."),
    jd: Optional[UploadFile] = File(None, description="The job description file (PDF or DOCX format). Optional if jd_text is provided."),
    jd_text: Optional[str] = Form(None, description="Plain text content of the job description (alternative to JD file upload).")
):
    """
    Matches a resume against a job description using the Google Gemini LLM.
    Calculates match score, identifies matched/missing skills, and suggests improvements.
    """
    logging.info(f"Received match request for resume: {resume.filename}")

    # 1. Extract text from the uploaded resume file
    resume_content_text = ""
    resume_bytes = await resume.read()
    if resume.filename.endswith(".pdf"):
        resume_content_text = extract_text_from_pdf(resume_bytes)
    elif resume.filename.lower().endswith((".docx", ".doc")): # Use .lower() for case-insensitivity
        resume_content_text = extract_text_from_docx(resume_bytes)
    else:
        logging.error(f"Unsupported resume file format: {resume.filename}")
        raise HTTPException(status_code=400, detail="Unsupported resume file format. Please upload a PDF or DOCX.")

    if not resume_content_text.strip(): # Check for truly empty text after extraction
        logging.error("Could not extract meaningful text from resume.")
        raise HTTPException(status_code=400, detail="Could not extract text from resume. Ensure it contains readable text.")

    # 2. Determine Job Description (JD) text source (file upload takes precedence over text form field)
    job_description_content_text = ""
    if jd:
        logging.info(f"Processing JD file: {jd.filename}")
        jd_bytes = await jd.read()
        if jd.filename.endswith(".pdf"):
            job_description_content_text = extract_text_from_pdf(jd_bytes)
        elif jd.filename.lower().endswith((".docx", ".doc")):
            job_description_content_text = extract_text_from_docx(jd_bytes)
        else:
            logging.error(f"Unsupported JD file format: {jd.filename}")
            raise HTTPException(status_code=400, detail="Unsupported JD file format. Please upload a PDF or DOCX.")
    elif jd_text: # If no file, use the provided text
        job_description_content_text = jd_text

    if not job_description_content_text.strip():
        logging.error("No job description provided or extracted.")
        raise HTTPException(status_code=400, detail="No job description provided or extracted. Please upload a JD file or provide JD text.")

    logging.info(f"Job Description (first 100 chars): {job_description_content_text[:100]}...")
    logging.info(f"Uploaded Resume Content (first 100 chars): {resume_content_text[:100]}...")

    # --- Use Gemini API for Matching and Suggestions ---
    prompt_text = f"""You are an expert resume analyzer. Your task is to compare a candidate's resume against a job description and provide a match analysis in JSON format.

**Resume:**
```
{resume_content_text}
```

**Job Description:**
```
{job_description_content_text}
```

Please provide the following information in a JSON object:
1.  **`match_score`**: An integer representing the overall percentage match of the resume to the job description (0-100). Be fair and realistic. This should be a single, numerical value.
2.  **`matched_skills`**: A list of key skills found in the resume that are also explicitly mentioned or clearly implied as required by the job description.
3.  **`missing_skills`**: A list of important skills mentioned in the job description that are either missing from the resume or not clearly highlighted. Prioritize up to 7 most critical missing skills.
4.  **`suggestions`**: A concise, actionable paragraph of 2-3 sentences advising the candidate on how to improve their resume to better match the job description, specifically addressing the missing skills and recommending how to integrate keywords naturally.

Return ONLY the JSON object. Do not include any conversational text, markdown outside the JSON, or code blocks (like ```json).

Example JSON Structure (ensure your output strictly follows this):
{{
  "match_score": 75,
  "matched_skills": ["Data Analysis", "Power BI", "SQL", "Dashboard Design"],
  "missing_skills": ["ETL Processes", "Data Modeling", "Azure"],
  "suggestions": "To enhance your resume, consider adding specific examples of your experience with ETL processes and data modeling. Elaborate on how you've used tools like Azure in past projects to demonstrate a broader technical skill set."
}}
"""
    
    payload = {
        "contents": [
            {
                "role": "user",
                "parts": [
                    {"text": prompt_text}
                ]
            }
        ],
        "generationConfig": {
            "temperature": 0.1, # Keep temperature low for more consistent results for matching
            "topK": 40,
            "topP": 0.95,
            "maxOutputTokens": 500, # Sufficient for structured output
        }
    }

    headers = {
        "Content-Type": "application/json"
    }
    params = {"key": GOOGLE_GEMINI_API_KEY}

    response_from_gemini = None # Initialize
    
    try:
        logging.info(f"Attempting to call Google Gemini API for match at: {GEMINI_API_BASE_URL}")
        
        # Defensive URL cleaning before request is sent
        cleaned_api_url = GEMINI_API_BASE_URL.replace('[', '').replace(']', '').replace('(', '').replace(')', '')
        logging.info(f"Using cleaned API URL for match: '{cleaned_api_url}'") 

        response_from_gemini = requests.post(
            cleaned_api_url,
            headers=headers,
            json=payload,
            params=params
        )
        response_from_gemini.raise_for_status() 

    except requests.exceptions.RequestException as e:
        logging.error(f"Google Gemini API request for match failed: {e}", exc_info=True)
        if response_from_gemini and response_from_gemini.text:
            logging.error(f"Gemini API Error Response for match: {response_from_gemini.text}")
        else:
            logging.error("Gemini API request for match failed before receiving any response.")
        raise HTTPException(status_code=500, detail=f"Failed to get match analysis from AI: {e}")

    try:
        gemini_response_data = response_from_gemini.json()
        generated_text = gemini_response_data["candidates"][0]["content"]["parts"][0]["text"]
        logging.info(f"Raw Gemini Match Response (first 500 chars): {generated_text[:500]}...")

        # Robust JSON extraction: Find the first '{' and the last '}'
        json_start = generated_text.find("{")
        json_end = generated_text.rfind("}")

        if json_start == -1 or json_end == -1 or json_end < json_start:
            logging.error(f"No complete JSON object found in Gemini match output. Generated text: {generated_text}")
            raise ValueError("No complete JSON object found in AI match output.")

        match_result_json_string = generated_text[json_start : json_end + 1]
        match_result_json_string = match_result_json_string.strip() # Remove whitespace

        match_data = json.loads(match_result_json_string)
        logging.info("Successfully parsed JSON from AI match response.")

        # Validate the structure of the AI's response for the match endpoint
        required_keys = ["match_score", "matched_skills", "missing_skills", "suggestions"]
        if not all(key in match_data for key in required_keys):
            logging.error(f"AI match response is missing required keys: {match_data}")
            raise ValueError("AI match response is malformed or missing required data.")

        # Ensure match_score is an integer
        if not isinstance(match_data.get("match_score"), int):
            try:
                match_data["match_score"] = int(match_data["match_score"])
            except (ValueError, TypeError):
                logging.warning(f"match_score is not an int, attempting conversion: {match_data.get('match_score')}")
                match_data["match_score"] = 0 # Default if conversion fails

        # Ensure lists are actual lists, default to empty if not
        if not isinstance(match_data.get("matched_skills"), list):
            match_data["matched_skills"] = []
        if not isinstance(match_data.get("missing_skills"), list):
            match_data["missing_skills"] = []
        if not isinstance(match_data.get("suggestions"), str):
            match_data["suggestions"] = "No specific suggestions provided by AI."


        return {
            "match_score": match_data["match_score"],
            "skill_match_score": match_data["match_score"], # Using overall match as skill match for now
            "text_similarity_score": match_data["match_score"], # Using overall match as text similarity for now
            "matched_skills": match_data["matched_skills"],
            "missing_skills": match_data["missing_skills"],
            "suggestions": match_data["suggestions"],
            "job_description_extracted": job_description_content_text # Still useful to return
        }

    except (json.JSONDecodeError, ValueError) as e:
        logging.error(f"Failed to parse JSON from AI match response. Error: {e}", exc_info=True)
        logging.error(f"Problematic AI Match Output (full):\\n{generated_text}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to parse AI match response. The AI might have generated malformed JSON. Error: {str(e)}. Full AI output (first 1000 chars): {generated_text[:1000]}"
        )
    except Exception as e:
        logging.error(f"An unexpected error occurred during AI match response processing: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred while processing AI match response: {e}")


@app.post("/generate_tailored_resume")
async def generate_tailored_resume(
    background_tasks: BackgroundTasks,
    resume_file: UploadFile = File(..., description="The candidate's resume file (PDF or DOCX format)."),
    jd_file: Optional[UploadFile] = File(None, description="The job description file (PDF or DOCX format). Optional if jd_text is provided."),
    jd_text_input: Optional[str] = Form(None, description="Plain text content of the job description (alternative to JD file upload)."),
    resume_style: str = Form("Chronological", description="Desired resume style: 'Chronological', 'Functional', or 'Hybrid'.")
):
    """
    Generates a job-targeted resume using the Google Gemini LLM based on an existing resume's content
    and a job description. Returns the generated resume as a DOCX file.
    """
    logging.info(f"Received tailored resume generation request for resume: {resume_file.filename}")

    # 1. Extract text from the uploaded resume file
    resume_bytes = await resume_file.read()
    if resume_file.filename.endswith(".pdf"):
        uploaded_resume_text = extract_text_from_pdf(resume_bytes)
    elif resume_file.filename.lower().endswith((".docx", ".doc")):
        uploaded_resume_text = extract_text_from_docx(resume_bytes)
    else:
        logging.error(f"Unsupported resume file format for tailored generation: {resume_file.filename}")
        raise HTTPException(status_code=400, detail="Unsupported resume file format. Please upload a PDF or DOCX.")

    if not uploaded_resume_text.strip():
        logging.error("Could not extract meaningful text from uploaded resume for tailored generation.")
        raise HTTPException(status_code=400, detail="Could not extract text from uploaded resume.")

    # 2. Determine Job Description (JD) text source
    job_description_text = jd_text_input
    if jd_file:
        logging.info(f"Processing JD file for tailored generation: {jd_file.filename}")
        jd_bytes = await jd.read()
        if jd.filename.endswith(".pdf"):
            job_description_text = extract_text_from_pdf(jd_bytes)
        elif jd.filename.lower().endswith((".docx", ".doc")):
            job_description_text = extract_text_from_docx(jd_bytes)
        else:
            logging.error(f"Unsupported JD file format for tailored generation: {jd_file.filename}")
            raise HTTPException(status_code=400, detail="Unsupported JD file format. Please upload a PDF or DOCX.")

    if not job_description_text.strip():
        logging.error("No job description provided or extracted for tailored generation.")
        raise HTTPException(status_code=400, detail="No job description provided or extracted.")

    logging.info(f"Job Description (first 100 chars): {job_description_text[:100]}...")
    logging.info(f"Uploaded Resume Content (first 100 chars): {uploaded_resume_text[:100]}...")

    # Define resume style guidelines based on the input
    style_guidelines = ""
    if resume_style.lower() == "chronological":
        style_guidelines = """
        Format the 'experience' section in reverse chronological order, clearly listing dates for each role. Emphasize career progression.
        """
    elif resume_style.lower() == "functional":
        style_guidelines = """
        Prioritize a 'Skills' or 'Summary of Qualifications' section at the top, detailing expertise areas and achievements. The 'experience' section can be brief, focusing on company names and dates, with detailed achievements moved to the skills/summary section.
        """
    elif resume_style.lower() == "hybrid":
        style_guidelines = """
        Combine a strong 'Summary of Qualifications' or 'Skills' section (like functional) with a 'Professional Experience' section that still details responsibilities and achievements in chronological order (like chronological).
        """
    else: # Default to chronological if invalid style is provided
        style_guidelines = """
        Format the 'experience' section in reverse chronological order, clearly listing dates for each role. Emphasize career progression.
        """
        logging.warning(f"Invalid resume style '{resume_style}' provided. Defaulting to Chronological.")

    # Craft the prompt for the LLM
    prompt_text = f"""You are a highly skilled professional resume writer. Your task is to create a new, job-targeted resume in JSON format.

Here is the **Candidate's Existing Resume**:
```
{uploaded_resume_text}
```

Here is the **Job Description** for which the candidate is applying:
```
{job_description_text}
```

Based on the **Candidate's Existing Resume**, extract all relevant information (e.g., name, contact, summary, skills, experience, education, certifications, projects). Then, **rewrite and tailor this information to best match the provided Job Description.**

**Specific Instructions:**
1.  **Achievements:** For each experience and project entry, quantify achievements with metrics and strong action verbs (e.g., "Increased sales by 15%", "Developed a system that reduced processing time by 20%"). Focus on impact and results.
2.  **Resume Style:** Adhere to a {resume_style.title()} resume style.
    {style_guidelines}
3.  **Sparse Resumes:** If the existing resume is very sparse or lacks detail, intelligently infer and generate plausible details (e.g., typical responsibilities, relevant projects, common skills) that align with the job description and the candidate's inferred background (based on name or any minimal info provided). Prioritize the most impactful and relevant information from the existing resume and supplement strategically.
4.  **Keyword Integration:** Identify key terms, skills, and phrases directly from the Job Description. Strategically and naturally integrate these keywords into the "summary," "skills," "experience" descriptions, and "project" descriptions to ensure high relevance and ATS compatibility. Do NOT just list them, weave them into the narrative and bullet points where appropriate.

Return only valid JSON in this exact structure, without any additional text, markdown formatting (like ```json), or conversational prose outside the JSON object. Ensure all fields are populated with extracted or generated information, or are empty strings/arrays if no relevant information can be inferred or generated.

{{
  "name": "Full Name",
  "summary": "Professional summary (3-4 lines, tailored to the job description, highlighting key strengths and career goals relevant to the JD)",
  "skills": ["Skill1", "Skill2", "SkillRelevantToJD", "Quantifiable Skill"],
  "experience": [
    {{"job_title": "Relevant Job Title", "company": "Company Name", "dates": "Start Date - End Date", "description": "Key responsibilities and **quantified achievements**, tailored to JD, using strong action verbs and integrated keywords"}}
  ],
  "education": [
    {{"degree": "Degree Name", "institution": "Institution Name", "year": "Year of Graduation"}}
  ],
  "certifications": ["Certification Name"],
  "projects": [
    {{"title": "Project Title", "tech_stack": ["Tech1", "Tech2"], "description": "Project description, highlighting relevance to JD and **quantified impact**, with integrated keywords"}}
  ],
  "contact": {{"email": "email@example.com", "phone": "123-456-7890", "linkedin": "https://linkedin.com/in/username"}}
}}
"""
    # Gemini API payload structure
    payload = {
        "contents": [
            {
                "role": "user",
                "parts": [
                    {"text": prompt_text}
                ]
            }
        ],
        "generationConfig": {
            "temperature": 0.7,
            "topK": 40,
            "topP": 0.95,
            "maxOutputTokens": 1800, # Increased max output tokens further to allow for more detailed resumes
        }
    }

    headers = {
        "Content-Type": "application/json"
    }
    params = {"key": GOOGLE_GEMINI_API_KEY}

    response = None 
    
    try:
        logging.info(f"Attempting to call Google Gemini API for generation at: {GEMINI_API_BASE_URL}")
        
        # Defensive URL cleaning before request is sent
        cleaned_api_url = GEMINI_API_BASE_URL.replace('[', '').replace(']', '').replace('(', '').replace(')', '')
        logging.info(f"Using cleaned API URL for generation: '{cleaned_api_url}'") 

        response = requests.post(
            cleaned_api_url, # Use the cleaned URL here
            headers=headers,
            json=payload,
            params=params # Pass API key as a query parameter
        )
        response.raise_for_status() # Raises an HTTPError for bad responses (4xx or 5xx)
    except requests.exceptions.RequestException as e:
        logging.error(f"Google Gemini API request for generation failed: {e}", exc_info=True)
        if response and response.text:
            logging.error(f"Gemini API Error Response for generation: {response.text}")
        elif response:
            logging.error(f"Gemini API Error Response for generation: (no text content available) Status Code: {response.status_code}")
        else:
            logging.error("Gemini API request for generation failed before receiving any response.")
        raise HTTPException(status_code=500, detail=f"Failed to generate tailored resume from AI: {e}")

    # 2. Parse the AI's response and extract JSON
    try:
        gemini_response_data = response.json()
        
        # Extract the generated text from Gemini's response structure
        if not gemini_response_data or "candidates" not in gemini_response_data or \
           not gemini_response_data["candidates"] or "content" not in gemini_response_data["candidates"][0] or \
           "parts" not in gemini_response_data["candidates"][0]["content"] or \
           not gemini_response_data["candidates"][0]["content"]["parts"] or \
           "text" not in gemini_response_data["candidates"][0]["content"]["parts"][0]:
            logging.error(f"Unexpected response structure from Google Gemini API: {gemini_response_data}")
            raise ValueError("Unexpected response structure from Google Gemini API.")

        generated_text = gemini_response_data["candidates"][0]["content"]["parts"][0]["text"]
        logging.info(f"Raw Gemini Generated Text (first 500 chars): {generated_text[:500]}...")

        # Robust JSON extraction: Find the first '{' and the last '}'
        json_start = generated_text.find("{")
        json_end = generated_text.rfind("}")

        if json_start == -1 or json_end == -1 or json_end < json_start:
            logging.error(f"No complete JSON object found in the AI output. Generated text: {generated_text}")
            raise ValueError("No complete JSON object found in the AI output.")

        resume_json_string = generated_text[json_start : json_end + 1]
        resume_json_string = resume_json_string.strip() # Remove leading/trailing whitespace

        data = json.loads(resume_json_string) # Parse the extracted JSON string into a Python dictionary
        logging.info("Successfully parsed JSON from AI response.")

    except (json.JSONDecodeError, ValueError) as e:
        logging.error(f"Failed to parse JSON from AI response. Error: {e}", exc_info=True)
        logging.error(f"Problematic AI Output (full):\\n{generated_text}") # Log the full problematic output
        raise HTTPException(
            status_code=500,
            detail=f"Failed to parse JSON from AI response. The AI might have generated malformed JSON. Error: {str(e)}. Full AI output (first 1000 chars): {generated_text[:1000]}"
        )
    except Exception as e:
        logging.error(f"An unexpected error occurred during AI response processing: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred while processing AI response: {e}")

    # 3. Generate DOCX from parsed JSON data
    doc = generate_resume_from_json(data)
    temp_filepath = ""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_f:
            temp_filepath = temp_f.name
            doc.save(temp_f.name) # Save the generated document
        
        # Add a background task to delete the temporary file after the response is sent
        background_tasks.add_task(os.remove, temp_filepath)

        logging.info(f"Generated AI resume DOCX at: {temp_filepath}")
        return FileResponse(
            temp_filepath,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="AI_Tailored_Resume.docx" # Changed filename to indicate it's tailored
        )
    except Exception as e:
        logging.error(f"Error saving generated AI resume DOCX: {e}", exc_info=True)
        if os.path.exists(temp_filepath):
            os.remove(temp_filepath) # Ensure cleanup even if saving fails
        raise HTTPException(status_code=500, detail=f"Failed to generate AI resume DOCX: {e}")




port = int(os.getenv("PORT", 8000))
# --- Main Entry Point ---
# Runs the FastAPI application using Uvicorn.
if __name__ == "__main__":
    # Host '0.0.0.0' makes the server accessible from other devices on the network,
    # not just localhost, which is useful for testing in different environments.
    # Port 8000 is the standard port used in your original setup.
    uvicorn.run(app, host="0.0.0.0", port=port)